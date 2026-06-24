"""
Generate a product catalog DOCX with SKU, title, use case, and applications.
Data sourced entirely from products.db — no API calls.
"""
import sys; sys.stdout.reconfigure(encoding='utf-8')
import sqlite3, json, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from rich.console import Console
from rich.progress import track

console = Console()
DB_PATH = "./products.db"
OUT_PATH = r"C:\Users\eugen\Downloads\BZB_Gear_Product_Catalog.docx"

# ─── Category display names and descriptions ──────────────────────────────────

CATEGORY_INFO = {
    "switcher":              ("🔀 Matrix & Video Switchers",      "Route multiple video sources to multiple displays. Core of any multi-screen AV installation."),
    "extender":              ("📡 Signal Extenders",              "Extend HDMI/DisplayPort signals over Cat cable, fiber, or wireless — beyond the 5m HDMI limit."),
    "distribution_amp":      ("🔁 Splitters & Distribution Amps", "Send one video source to multiple displays simultaneously."),
    "av_over_ip":            ("🌐 AV Over IP",                   "Distribute video across a standard IP network — unlimited distance, scalable to hundreds of endpoints."),
    "camera":                ("📷 Cameras",                       "PTZ, ePTZ, box cameras, and webcams for conferencing, broadcast, streaming, and production."),
    "kvm_switch":            ("⌨️  KVM Switches",                 "Control multiple computers from one keyboard, monitor, and mouse. Often includes video switching."),
    "multiviewer":           ("🖥️  Multiviewers & Video Walls",   "Display multiple video sources on a single screen, or drive large video wall arrays."),
    "presentation_switcher": ("📊 Presentation Switchers",        "Connect laptops and devices to a room display — often with wireless/BYOD support."),
    "encoder_decoder":       ("📦 Encoders & Decoders",           "Convert HDMI to IP streams (encoders) or IP streams back to HDMI (decoders). Used for streaming and AV-over-IP."),
    "capture":               ("🎬 Capture Cards & Converters",    "Capture video from HDMI/SDI sources into a computer for recording, streaming, or processing."),
    "sdi":                   ("📺 SDI Equipment & Converters",    "Professional broadcast-grade SDI signals — converters, distribution amps, and splitters."),
    "audio":                 ("🔊 Audio Equipment",               "Amplifiers, DSP processors, speakerphones, microphones, and audio converters."),
    "controller":            ("🕹️  Joystick Controllers",          "Hardware controllers for operating PTZ cameras — with preview screens and multi-camera support."),
    "signal_generator":      ("⚡ Signal Generators",             "Generate test patterns and HDMI/SDI signals for system testing, calibration, and troubleshooting."),
    "network":               ("🔌 Network Switches",              "Managed and unmanaged PoE switches — backbone for AV-over-IP and powered camera installations."),
    "integration":           ("🔧 Integration Tools",             "Control systems, scalers, and tools for integrating AV equipment into automation systems."),
    "usb_extender":          ("🖱️  USB Extenders",                "Extend USB signals over Cat cable — for keyboards, mice, and USB peripherals at a distance."),
    "videobar":              ("📹 Videobars",                     "All-in-one conference room cameras with integrated microphone and speaker."),
    "medical_cart":          ("🏥 Medical Carts",                 "Mobile medical-grade display carts for clinical environments."),
    "cable_kit":             ("🔗 Cables & Fiber Kits",           "HDMI, SDI, fiber optic cables, and pre-terminated fiber extension kits."),
    "accessory":             ("🔩 Accessories & Mounts",          "Mounting brackets, SFP modules, adapters, and other accessories."),
    "bundle":                ("📦 Production Bundles",            "Complete camera system bundles — cameras + controller + switch + cables, ready to deploy."),
    "encoder_decoder":       ("📦 Encoders & Decoders",           "HDMI to IP stream encoders and IP to HDMI decoders for streaming and AV-over-IP deployments."),
}

# ─── Use case synthesis from description + features ──────────────────────────

def synthesize_use_case(row: dict) -> tuple[str, str]:
    """Returns (use_case_summary, when_to_use) from product data."""
    title   = row.get("title") or ""
    desc    = row.get("description") or ""
    feats   = json.loads(row.get("features") or "[]")
    specs   = json.loads(row.get("specs_json") or "{}")
    cat     = row.get("category") or ""
    site_sub = row.get("site_subcategory") or ""

    # Build combined text for analysis
    feat_text = " | ".join(feats[:6])
    spec_vals = " | ".join(list(specs.values())[:8])

    # ── Use case: first 2 sentences of description, cleaned up ──
    desc_clean = re.sub(r'\s+', ' ', desc).strip()
    sentences = re.split(r'(?<=[.!?])\s+', desc_clean)
    use_case = " ".join(sentences[:2]).strip()
    if len(use_case) > 400:
        use_case = use_case[:400].rsplit(" ", 1)[0] + "…"
    if not use_case:
        use_case = title

    # ── When to apply: infer from category + title keywords ──
    when = []
    t = title.lower()
    d = (desc + feat_text).lower()

    # Distance/cable
    dist = row.get("max_distance_m")
    if dist:
        when.append(f"Runs up to {dist}m — use when displays are far from source")

    # Resolution
    res = json.loads(row.get("resolutions") or "[]")
    if "8K60" in res or "8K30" in res:
        when.append("8K installations requiring future-proof signal chain")
    elif "4K120" in res:
        when.append("High-refresh gaming or simulation rooms needing 4K120Hz")
    elif "4K60" in res:
        when.append("4K60 professional AV installations")

    # Port counts
    inp = row.get("inputs")
    out = row.get("outputs")
    if inp and out and inp > 1 and out > 1:
        when.append(f"Multi-source installations: {inp} sources to {out} displays")
    elif inp == 1 and out and out > 1:
        when.append(f"Single source split to {out} displays")

    # Category-based when-to-use
    cat_whens = {
        "switcher":          "Conference rooms, sport bars, control rooms, broadcast studios needing flexible routing",
        "extender":          "Long cable runs where HDMI won't reach — retail, education, large venues",
        "distribution_amp":  "Duplicate a source to multiple screens: digital signage, lobby displays",
        "av_over_ip":        "Large venues, campus-wide distribution, 100+ endpoints over existing network",
        "camera":            "Video conferencing, live streaming, lecture capture, broadcast production",
        "kvm_switch":        "Multi-PC workstations, server rooms, trading desks, broadcast control rooms",
        "multiviewer":       "Security monitoring walls, broadcast control rooms, multi-feed sports viewing",
        "presentation_switcher": "Meeting rooms, classrooms, BYOD environments",
        "encoder_decoder":   "Live streaming to YouTube/Twitch, IPTV distribution, remote production",
        "capture":           "Recording gameplay/presentations, live streaming from HDMI sources",
        "sdi":               "Broadcast facilities, live production, long SDI cable runs (up to 300m)",
        "audio":             "Conference rooms, lecture halls, installed audio systems",
        "controller":        "Operating PTZ cameras in live production, houses of worship, sports broadcasting",
        "signal_generator":  "AV system commissioning, display calibration, signal troubleshooting",
        "network":           "PoE backbone for IP cameras, AV-over-IP endpoints, powered AV devices",
        "integration":       "Crestron/AMX/Control4 integration, corporate AV automation",
        "usb_extender":      "Remote USB peripherals — keyboards, mice, USB cameras at a distance",
        "videobar":          "Small-to-medium conference rooms needing all-in-one video/audio",
        "medical_cart":      "Clinical environments requiring mobile display with medical-grade compliance",
        "cable_kit":         "Connecting AV equipment — choose by signal type and run length",
        "accessory":         "Supplementary hardware for installed AV systems",
        "bundle":            "Complete turnkey studio or production system — faster deployment",
    }

    if cat in cat_whens:
        when.append(cat_whens[cat])

    # Wireless / wireless extender
    if "wireless" in t or "wireless" in d[:200]:
        when.append("Temporary installations or spaces where cabling is impractical")

    # NDI
    if "ndi" in t.lower() or "ndi" in d[:300].lower():
        when.append("Software-defined production workflows using vMix, OBS, TriCaster")

    # PoE
    if "poe" in t.lower() or "poe" in feat_text.lower():
        when.append("Simplifies installation — power and signal over a single Cat cable")

    when_text = "; ".join(when[:4]) if when else "General AV installation"
    return use_case, when_text


# ─── Document styling helpers ─────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx: int, width_cm: float):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p

def para_format(p, size_pt: int = 10, bold: bool = False, color=None, space_after: int = 0):
    run = p.runs[0] if p.runs else p.add_run()
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(space_after)


# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row

    # Load all products
    rows = conn.execute("""
        SELECT id, name, title, category, site_category, site_subcategory,
               description, features, specs_json, resolutions,
               inputs, outputs, max_distance_m, max_bandwidth_gbps,
               price_usd, price_sale_usd, stock_status,
               product_url, manual_url, image_url
        FROM products
        ORDER BY category, id
    """).fetchall()
    conn.close()

    console.print(f"Generating catalog for [bold]{len(rows)}[/bold] products...")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for level in [1, 2, 3]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x1a, 0x37, 0x6c)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BZB Gear Product Catalog")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x37, 0x6c)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("SKU · Title · Use Case · When to Apply")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"{len(rows)} products across 22 categories  |  June 2026")
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

    doc.add_page_break()

    # Group by category
    from collections import defaultdict
    by_cat: dict[str, list] = defaultdict(list)
    for row in rows:
        by_cat[dict(row)["category"]].append(dict(row))

    # Column headers + widths (cm): SKU | Title | Use Case | When to Apply | Price | Status
    COL_WIDTHS = [3.2, 5.5, 7.5, 5.5, 1.8, 2.0]
    HEADERS    = ["SKU", "Title", "What it does / Use Case", "When to apply", "Price", "Status"]
    HDR_COLOR  = "1A376C"   # dark blue
    ROW_COLOR1 = "FFFFFF"
    ROW_COLOR2 = "F0F4FA"

    for cat_key in sorted(by_cat.keys()):
        products = by_cat[cat_key]
        cat_label, cat_desc = CATEGORY_INFO.get(cat_key, (cat_key.replace("_"," ").title(), ""))

        # Category heading
        h = doc.add_heading(cat_label, level=1)

        p = doc.add_paragraph(cat_desc)
        p.paragraph_format.space_after = Pt(6)
        pr = p.runs[0] if p.runs else p.add_run(cat_desc)
        pr.font.size = Pt(10)
        pr.font.italic = True
        pr.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

        # Table
        table = doc.add_table(rows=1, cols=len(HEADERS))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        hrow = table.rows[0]
        for i, (hdr, w) in enumerate(zip(HEADERS, COL_WIDTHS)):
            cell = hrow.cells[i]
            cell.width = Cm(w)
            set_cell_bg(cell, HDR_COLOR)
            p = cell.paragraphs[0]
            run = p.add_run(hdr)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.paragraph_format.space_after = Pt(0)

        # Product rows
        for i, prod in enumerate(track(products, description=f"  {cat_label[:30]}...", total=len(products))):
            use_case, when = synthesize_use_case(prod)

            row = table.add_row()
            bg = ROW_COLOR2 if i % 2 else ROW_COLOR1

            vals = [
                prod["id"] or "",
                prod["title"] or prod["name"] or "",
                use_case,
                when,
                f"${prod['price_usd']:,.0f}" if prod.get("price_usd") else "—",
                prod.get("stock_status") or "—",
            ]

            for j, (val, w) in enumerate(zip(vals, COL_WIDTHS)):
                cell = row.cells[j]
                cell.width = Cm(w)
                set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                run = p.add_run(str(val)[:500])
                run.font.size = Pt(8.5 if j > 1 else 9)
                run.font.bold = (j == 0)  # bold SKU
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(1)

        doc.add_paragraph()  # spacing between categories

    doc.save(OUT_PATH)
    console.print(f"\n[bold green]Saved: {OUT_PATH}[/bold green]")
    console.print(f"Size: {Path(OUT_PATH).stat().st_size // 1024} KB")


if __name__ == "__main__":
    main()
