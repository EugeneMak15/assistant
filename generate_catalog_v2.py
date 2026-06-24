"""
Generate BZB Gear product catalog DOCX v2.
Uses AI-generated what_it_does + use_cases fields.
Columns: SKU | Title | What it does | Use Cases | Price | Status
"""
import sys; sys.stdout.reconfigure(encoding='utf-8')
import sqlite3, json
from pathlib import Path
from collections import defaultdict

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from rich.console import Console

console = Console()
DB_PATH   = "./products.db"
OUT_PATH  = r"C:\Users\eugen\Downloads\BZB_Gear_Product_Catalog.docx"

# ─── Category display names ───────────────────────────────────────────────────

CATEGORIES = {
    "switcher":              "Matrix & Video Switchers",
    "extender":              "Signal Extenders",
    "distribution_amp":      "Splitters & Distribution Amps",
    "av_over_ip":            "AV Over IP",
    "camera":                "Cameras",
    "kvm_switch":            "KVM Switches",
    "multiviewer":           "Multiviewers & Video Walls",
    "presentation_switcher": "Presentation Switchers",
    "encoder_decoder":       "Encoders & Decoders",
    "capture":               "Capture Cards & Converters",
    "sdi":                   "SDI Equipment",
    "audio":                 "Audio Equipment",
    "controller":            "PTZ Joystick Controllers",
    "signal_generator":      "Signal Generators",
    "network":               "Network Switches",
    "integration":           "Integration Tools",
    "usb_extender":          "USB Extenders",
    "videobar":              "Videobars",
    "medical_cart":          "Medical Carts",
    "cable_kit":             "Cables & Fiber Kits",
    "accessory":             "Accessories & Mounts",
    "bundle":                "Production Bundles",
}

CAT_ORDER = list(CATEGORIES.keys())

# ─── Colour palette ───────────────────────────────────────────────────────────

BLUE_DARK  = "1A376C"
BLUE_MID   = "2E5FA3"
BLUE_LIGHT = "D6E4F7"
WHITE      = "FFFFFF"
GREY_LIGHT = "F4F6FA"
GREY_MID   = "E2E8F4"

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cell_text(cell, text: str, size: int = 9, bold: bool = False,
              color: tuple = None, italic: bool = False, align=None):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if align:
        p.alignment = align
    run = p.add_run(str(text)[:800])
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_bullet_cell(cell, items: list[str], size: int = 8):
    """Fill a cell with bullet-list items, one paragraph per item."""
    first = True
    for item in items[:10]:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(f"• {item[:200]}")
        run.font.size = Pt(size)

def set_row_height(row, height_cm: float):
    trPr = row._tr.get_or_add_trPr()
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(height_cm * 567)))  # 567 twips per cm
    trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row

    rows = conn.execute("""
        SELECT id, name, title, category, site_category, site_subcategory,
               inputs, outputs, max_distance_m, max_bandwidth_gbps,
               resolutions, input_signals,
               price_usd, price_sale_usd, stock_status,
               what_it_does, use_cases
        FROM products
        ORDER BY category, id
    """).fetchall()
    conn.close()

    enriched = sum(1 for r in rows if r["what_it_does"])
    console.print(f"Products: {len(rows)}  |  Enriched: {enriched}  |  Pending: {len(rows)-enriched}")

    if enriched < len(rows):
        console.print("[yellow]Warning: some products not yet enriched. Run enrich_products.py first.[/yellow]")

    # Group by category
    by_cat = defaultdict(list)
    for row in rows:
        by_cat[dict(row)["category"]].append(dict(row))

    # ── Build document ──
    doc = Document()

    for sec in doc.sections:
        sec.page_width    = Cm(29.7)   # A4 landscape
        sec.page_height   = Cm(21.0)
        sec.top_margin    = Cm(1.2)
        sec.bottom_margin = Cm(1.2)
        sec.left_margin   = Cm(1.5)
        sec.right_margin  = Cm(1.5)

    # Default style
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Title page ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BZB Gear — Product Catalog")
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{len(rows)} products  ·  {len(by_cat)} categories  ·  June 2026")
    r2.font.size  = Pt(12)
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x99)

    doc.add_page_break()

    # Column layout (cm) — landscape A4 content width ~26.7cm
    # SKU | Title | What it does | Use Cases | Price | Status
    COL_W = [3.0, 5.0, 7.5, 8.5, 1.8, 1.9]
    HEADERS = ["SKU", "Title", "What It Does", "Use Cases", "Price", "Status"]

    for cat_key in CAT_ORDER:
        products = by_cat.get(cat_key, [])
        if not products:
            continue

        cat_name = CATEGORIES.get(cat_key, cat_key.replace("_", " ").title())

        # Category heading paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(cat_name)
        run.font.size  = Pt(16)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)

        # Sub-line: product count + site category
        site_cats = sorted(set(
            p.get("site_category") for p in products if p.get("site_category")
        ))
        if site_cats:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after  = Pt(6)
            r2 = p2.add_run(f"{len(products)} products  ·  {' / '.join(site_cats)}")
            r2.font.size   = Pt(9)
            r2.font.italic = True
            r2.font.color.rgb = RGBColor(0x66, 0x77, 0x99)

        # Table
        table = doc.add_table(rows=1, cols=len(HEADERS))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row
        hrow = table.rows[0]
        set_row_height(hrow, 0.7)
        for i, (hdr, w) in enumerate(zip(HEADERS, COL_W)):
            cell = hrow.cells[i]
            cell.width = Cm(w)
            set_cell_bg(cell, BLUE_DARK)
            cell_text(cell, hdr, size=9, bold=True,
                      color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)

        # Product rows
        for idx, prod in enumerate(products):
            bg = GREY_LIGHT if idx % 2 == 0 else WHITE

            # Specs line for title column
            specs_bits = []
            if prod.get("inputs") and prod.get("outputs"):
                specs_bits.append(f"{prod['inputs']}×{prod['outputs']}")
            if prod.get("max_distance_m"):
                specs_bits.append(f"{prod['max_distance_m']}m")
            res = json.loads(prod.get("resolutions") or "[]")
            if res:
                specs_bits.append(res[-1])  # highest resolution
            bw = prod.get("max_bandwidth_gbps")
            if bw:
                specs_bits.append(f"{bw}Gbps")
            specs_str = "  ·  ".join(specs_bits)

            # Use cases
            use_cases_raw = prod.get("use_cases") or "[]"
            try:
                use_cases = json.loads(use_cases_raw)
                if not isinstance(use_cases, list):
                    use_cases = [str(use_cases)]
            except Exception:
                use_cases = [use_cases_raw[:200]]

            # What it does
            what = prod.get("what_it_does") or ""
            if not what:
                what = "(not yet enriched — run enrich_products.py)"

            # Price display
            price = prod.get("price_usd")
            sale  = prod.get("price_sale_usd")
            if sale and sale < (price or 0):
                price_str = f"${sale:,.0f}\n(was ${price:,.0f})"
            elif price:
                price_str = f"${price:,.0f}"
            else:
                price_str = "—"

            row = table.add_row()
            set_row_height(row, 1.0)

            vals_simple = [
                (prod["id"],    9,  True,  False, None),
                (None,          9,  False, False, None),   # title handled separately
                (what,          8,  False, False, None),
                (None,          8,  False, False, None),   # use_cases handled separately
                (price_str,     8,  False, False, WD_ALIGN_PARAGRAPH.CENTER),
                (prod.get("stock_status") or "—", 8, False, False, WD_ALIGN_PARAGRAPH.CENTER),
            ]

            for j, (w_cm) in enumerate(COL_W):
                cell = row.cells[j]
                cell.width = Cm(w_cm)
                set_cell_bg(cell, bg)

            # SKU
            cell_text(row.cells[0], prod["id"],
                      size=9, bold=True, color=(0x1A, 0x37, 0x6C))

            # Title + specs
            tc = row.cells[1]
            p_title = tc.paragraphs[0]
            p_title.paragraph_format.space_before = Pt(1)
            p_title.paragraph_format.space_after  = Pt(1)
            r_title = p_title.add_run(prod.get("title") or prod.get("name") or "")
            r_title.font.size = Pt(8.5)
            r_title.font.bold = True
            if specs_str:
                p_spec = tc.add_paragraph()
                p_spec.paragraph_format.space_before = Pt(0)
                p_spec.paragraph_format.space_after  = Pt(0)
                r_spec = p_spec.add_run(specs_str)
                r_spec.font.size  = Pt(7.5)
                r_spec.font.italic = True
                r_spec.font.color.rgb = RGBColor(0x55, 0x66, 0x99)

            # What it does
            cell_text(row.cells[2], what, size=8)

            # Use cases — bullet list
            add_bullet_cell(row.cells[3], use_cases, size=8)

            # Price
            cell_text(row.cells[4], price_str, size=8,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

            # Status
            status = prod.get("stock_status") or "—"
            status_color = (0x1A, 0x7A, 0x3A) if "stock" in status.lower() \
                      else (0xC0, 0x60, 0x10) if "pre" in status.lower() \
                      else (0x55, 0x55, 0x55)
            cell_text(row.cells[5], status, size=8,
                      color=status_color, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Spacing after table
        doc.add_paragraph()

    doc.save(OUT_PATH)
    size_kb = Path(OUT_PATH).stat().st_size // 1024
    console.print(f"\n[bold green]Saved: {OUT_PATH}[/bold green]  ({size_kb} KB)")


if __name__ == "__main__":
    main()
